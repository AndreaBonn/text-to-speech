"""
tests/test_converters.py
Test per il modulo converters.py: dispatcher e convertitori di formato.

Ogni convertitore è testato con file reali creati in tmp_path.
"""

from pathlib import Path
from unittest.mock import patch

import pytest

from converters import SUPPORTED_EXTENSIONS, file_a_testo

# ===========================================================================
# Test — Dispatcher file_a_testo
# ===========================================================================


class TestFileATesto:
    """Test per il dispatcher principale."""

    def test_estensione_non_supportata_solleva_errore(self, tmp_path):
        """Un file con estensione non supportata deve sollevare ValueError."""
        # Arrange
        file_csv = tmp_path / "dati.csv"
        file_csv.write_text("a,b,c")

        # Act & Assert
        with pytest.raises(ValueError, match="non supportato"):
            file_a_testo(file_csv)

    def test_supported_extensions_contiene_tutti_i_formati(self):
        """SUPPORTED_EXTENSIONS deve contenere tutti i formati dichiarati."""
        # Assert
        attesi = {".md", ".txt", ".epub", ".docx", ".html", ".htm", ".pdf"}
        assert attesi == SUPPORTED_EXTENSIONS

    def test_dispatcher_case_insensitive(self, tmp_path):
        """L'estensione deve essere case-insensitive."""
        # Arrange
        file_txt = tmp_path / "test.TXT"
        file_txt.write_text("Contenuto del file.")

        # Act
        risultato = file_a_testo(file_txt)

        # Assert
        assert risultato == "Contenuto del file."


# ===========================================================================
# Test — Convertitore .txt
# ===========================================================================


class TestConvertiTesto:
    """Test per il convertitore di file testo puro."""

    def test_legge_testo_semplice(self, tmp_path):
        """Un file .txt deve essere letto e restituito invariato."""
        # Arrange
        txt = tmp_path / "nota.txt"
        txt.write_text("Questa è una nota semplice.")

        # Act
        risultato = file_a_testo(txt)

        # Assert
        assert risultato == "Questa è una nota semplice."

    def test_strip_whitespace(self, tmp_path):
        """Spazi iniziali e finali devono essere rimossi."""
        # Arrange
        txt = tmp_path / "spazi.txt"
        txt.write_text("  \n\nContenuto con spazi.\n\n  ")

        # Act
        risultato = file_a_testo(txt)

        # Assert
        assert risultato == "Contenuto con spazi."

    def test_file_vuoto(self, tmp_path):
        """Un file vuoto deve restituire stringa vuota."""
        # Arrange
        txt = tmp_path / "vuoto.txt"
        txt.write_text("")

        # Act
        risultato = file_a_testo(txt)

        # Assert
        assert risultato == ""

    def test_unicode_preservato(self, tmp_path):
        """Caratteri unicode (accenti, emoji) devono essere preservati."""
        # Arrange
        txt = tmp_path / "unicode.txt"
        txt.write_text("Caffè, più, naïve, 日本語")

        # Act
        risultato = file_a_testo(txt)

        # Assert
        assert "Caffè" in risultato
        assert "日本語" in risultato

    def test_paragrafi_multipli(self, tmp_path):
        """Più paragrafi separati da righe vuote devono essere preservati."""
        # Arrange
        txt = tmp_path / "multi.txt"
        txt.write_text("Primo paragrafo.\n\nSecondo paragrafo.\n\nTerzo.")

        # Act
        risultato = file_a_testo(txt)

        # Assert
        assert "Primo paragrafo." in risultato
        assert "Secondo paragrafo." in risultato
        assert "Terzo." in risultato


# ===========================================================================
# Test — Convertitore .md (delega a leggi)
# ===========================================================================


class TestConvertiMarkdown:
    """Test per il convertitore Markdown (delega a markdown_a_testo)."""

    def test_rimuove_header_markdown(self, tmp_path):
        """I titoli # devono essere rimossi dal testo."""
        # Arrange
        md = tmp_path / "doc.md"
        md.write_text("# Titolo\n\nContenuto del documento.")

        # Act — forza fallback regex (senza pandoc)
        with patch("converters.shutil.which", return_value=None):
            risultato = file_a_testo(md)

        # Assert
        assert "#" not in risultato
        assert "Titolo" in risultato
        assert "Contenuto del documento." in risultato


# ===========================================================================
# Test — Convertitore .docx
# ===========================================================================


class TestConvertiDocx:
    """Test per il convertitore Word DOCX."""

    def test_estrae_paragrafi(self, tmp_path):
        """Un DOCX con paragrafi deve restituire il testo separato."""
        from docx import Document

        # Arrange
        docx_path = tmp_path / "documento.docx"
        doc = Document()
        doc.add_paragraph("Primo paragrafo del documento.")
        doc.add_paragraph("Secondo paragrafo con contenuto.")
        doc.save(str(docx_path))

        # Act
        risultato = file_a_testo(docx_path)

        # Assert
        assert "Primo paragrafo del documento." in risultato
        assert "Secondo paragrafo con contenuto." in risultato

    def test_ignora_paragrafi_vuoti(self, tmp_path):
        """Paragrafi vuoti nel DOCX devono essere ignorati."""
        from docx import Document

        # Arrange
        docx_path = tmp_path / "vuoti.docx"
        doc = Document()
        doc.add_paragraph("Testo valido.")
        doc.add_paragraph("")  # vuoto
        doc.add_paragraph("   ")  # solo spazi
        doc.add_paragraph("Altro testo.")
        doc.save(str(docx_path))

        # Act
        risultato = file_a_testo(docx_path)
        paragrafi = [p for p in risultato.split("\n\n") if p.strip()]

        # Assert
        assert len(paragrafi) == 2

    def test_docx_vuoto(self, tmp_path):
        """Un DOCX senza contenuto deve restituire stringa vuota."""
        from docx import Document

        # Arrange
        docx_path = tmp_path / "empty.docx"
        doc = Document()
        doc.save(str(docx_path))

        # Act
        risultato = file_a_testo(docx_path)

        # Assert
        assert risultato == ""


# ===========================================================================
# Test — Convertitore .html
# ===========================================================================


class TestConvertiHtml:
    """Test per il convertitore HTML."""

    def test_estrae_testo_body(self, tmp_path):
        """Il testo del body deve essere estratto."""
        # Arrange
        html = tmp_path / "pagina.html"
        html.write_text("""<!DOCTYPE html>
<html><head><title>Test</title></head>
<body><h1>Titolo</h1><p>Contenuto della pagina.</p></body>
</html>""")

        # Act
        risultato = file_a_testo(html)

        # Assert
        assert "Titolo" in risultato
        assert "Contenuto della pagina." in risultato

    def test_rimuove_script_e_style(self, tmp_path):
        """Tag script e style devono essere rimossi."""
        # Arrange
        html = tmp_path / "scripts.html"
        html.write_text("""<html><body>
<script>alert('xss')</script>
<style>body { color: red; }</style>
<p>Testo visibile.</p>
</body></html>""")

        # Act
        risultato = file_a_testo(html)

        # Assert
        assert "alert" not in risultato
        assert "color: red" not in risultato
        assert "Testo visibile." in risultato

    def test_rimuove_nav_header_footer(self, tmp_path):
        """Elementi di navigazione devono essere rimossi."""
        # Arrange
        html = tmp_path / "layout.html"
        html.write_text("""<html><body>
<nav><a href="/">Home</a><a href="/about">About</a></nav>
<header><h1>Header del sito</h1></header>
<main><p>Contenuto principale.</p></main>
<footer><p>Copyright 2025</p></footer>
</body></html>""")

        # Act
        risultato = file_a_testo(html)

        # Assert
        assert "Contenuto principale." in risultato
        assert "Home" not in risultato
        assert "Copyright" not in risultato

    def test_preferisce_main_content(self, tmp_path):
        """Se presente, deve estrarre da <main> o <article>."""
        # Arrange
        html = tmp_path / "article.html"
        html.write_text("""<html><body>
<div class="sidebar">Menu laterale con molto testo.</div>
<main><p>Articolo importante.</p></main>
</body></html>""")

        # Act
        risultato = file_a_testo(html)

        # Assert
        assert "Articolo importante." in risultato

    def test_html_estensione_htm(self, tmp_path):
        """Anche .htm deve funzionare come .html."""
        # Arrange
        htm = tmp_path / "pagina.htm"
        htm.write_text("<html><body><p>Testo HTM.</p></body></html>")

        # Act
        risultato = file_a_testo(htm)

        # Assert
        assert "Testo HTM." in risultato


# ===========================================================================
# Test — Convertitore .pdf
# ===========================================================================


class TestConvertiPdf:
    """Test per il convertitore PDF."""

    def test_estrae_testo_da_pdf(self, tmp_path):
        """Un PDF con testo deve essere estratto correttamente."""
        import pymupdf

        # Arrange — crea un PDF di test con pymupdf
        pdf_path = tmp_path / "documento.pdf"
        doc = pymupdf.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Primo paragrafo del documento PDF.")
        page.insert_text((72, 120), "Secondo paragrafo con contenuto.")
        doc.save(str(pdf_path))
        doc.close()

        # Act
        risultato = file_a_testo(pdf_path)

        # Assert
        assert "Primo paragrafo del documento PDF." in risultato
        assert "Secondo paragrafo con contenuto." in risultato

    def test_pdf_multipagina(self, tmp_path):
        """Un PDF con più pagine deve estrarre testo da tutte."""
        import pymupdf

        # Arrange
        pdf_path = tmp_path / "multi.pdf"
        doc = pymupdf.open()
        for i in range(3):
            page = doc.new_page()
            page.insert_text((72, 72), f"Contenuto pagina {i + 1}.")
        doc.save(str(pdf_path))
        doc.close()

        # Act
        risultato = file_a_testo(pdf_path)

        # Assert
        assert "Contenuto pagina 1." in risultato
        assert "Contenuto pagina 2." in risultato
        assert "Contenuto pagina 3." in risultato

    def test_rimuove_numeri_pagina_isolati(self, tmp_path):
        """Numeri di pagina isolati su una riga devono essere rimossi."""
        import pymupdf

        # Arrange
        pdf_path = tmp_path / "paginato.pdf"
        doc = pymupdf.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Testo del documento.")
        page.insert_text((300, 780), "1")  # numero di pagina in basso
        doc.save(str(pdf_path))
        doc.close()

        # Act
        risultato = file_a_testo(pdf_path)

        # Assert
        assert "Testo del documento." in risultato
        # Il numero "1" isolato deve essere rimosso
        lines = [line.strip() for line in risultato.split("\n") if line.strip()]
        assert "1" not in lines


# ===========================================================================
# Test — Convertitore .epub
# ===========================================================================


class TestConvertiEpub:
    """Test per il convertitore EPUB."""

    def _make_epub(self, tmp_path: Path, capitoli: list[str]) -> Path:
        """Helper: crea un EPUB minimale con i capitoli dati."""
        from ebooklib import epub

        book = epub.EpubBook()
        book.set_identifier("test-id-123")
        book.set_title("Test Book")
        book.set_language("it")

        items = []
        for i, testo in enumerate(capitoli):
            ch = epub.EpubHtml(
                title=f"Capitolo {i + 1}",
                file_name=f"chap_{i}.xhtml",
                lang="it",
            )
            ch.content = f"<html><body><p>{testo}</p></body></html>".encode()
            book.add_item(ch)
            items.append(ch)

        book.toc = items
        book.spine = ["nav", *items]
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())

        epub_path = tmp_path / "libro.epub"
        epub.write_epub(str(epub_path), book)
        return epub_path

    def test_estrae_testo_capitoli(self, tmp_path):
        """Un EPUB con capitoli deve estrarre il testo di ciascuno."""
        # Arrange
        epub_path = self._make_epub(
            tmp_path,
            ["Contenuto del primo capitolo.", "Contenuto del secondo capitolo."],
        )

        # Act
        risultato = file_a_testo(epub_path)

        # Assert
        assert "Contenuto del primo capitolo." in risultato
        assert "Contenuto del secondo capitolo." in risultato

    def test_epub_singolo_capitolo(self, tmp_path):
        """Un EPUB con un solo capitolo deve funzionare."""
        # Arrange
        epub_path = self._make_epub(tmp_path, ["Unico capitolo del libro."])

        # Act
        risultato = file_a_testo(epub_path)

        # Assert
        assert "Unico capitolo del libro." in risultato

    def test_epub_rimuove_script(self, tmp_path):
        """Eventuali tag script nell'EPUB devono essere rimossi."""
        from ebooklib import epub

        # Arrange
        book = epub.EpubBook()
        book.set_identifier("test-script")
        book.set_title("Script Test")
        book.set_language("it")

        ch = epub.EpubHtml(title="Cap", file_name="ch.xhtml", lang="it")
        ch.content = b"""<html><body>
<script>alert('evil')</script>
<p>Testo sicuro.</p>
</body></html>"""
        book.add_item(ch)
        book.spine = ["nav", ch]
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())

        epub_path = tmp_path / "script.epub"
        epub.write_epub(str(epub_path), book)

        # Act
        risultato = file_a_testo(epub_path)

        # Assert
        assert "alert" not in risultato
        assert "Testo sicuro." in risultato
